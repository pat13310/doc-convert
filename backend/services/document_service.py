"""
Service de gestion des documents
Contient les fonctions pour extraire et convertir les documents
"""
import os
import tempfile
import logging
import docx
from docx import Document
import subprocess
import sys
import uuid
import zipfile
import fitz  # PyMuPDF
import pandas as pd
import csv
import chardet
import re
from docx.shared import Inches
import openpyxl
import xlrd
import traceback
import shutil
from io import StringIO
from PIL import Image
from pathlib import Path

# Configuration du logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path):
    """
    Extrait le texte d'un fichier DOCX
    """
    try:
        logger.info(f"Extraction du texte du fichier DOCX: {file_path}")
        doc = Document(file_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        return '\n'.join(text)
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte DOCX: {str(e)}")
        logger.error(traceback.format_exc())
        raise Exception(f"Erreur lors de l'extraction du texte DOCX: {str(e)}")

def extract_text_from_pdf(file_path):
    """
    Extrait le texte d'un fichier PDF en gérant les émoticônes et caractères spéciaux
    
    Args:
        file_path (str): Chemin vers le fichier PDF
        
    Returns:
        str: Texte extrait du fichier
        
    Raises:
        Exception: En cas d'erreur lors de l'extraction
    """
    logger = logging.getLogger(__name__)
    try:
        logger.info(f"Extraction du texte du fichier PDF: {file_path}")
        
        # Utiliser PyMuPDF (fitz) pour extraire le texte
        text = ""
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text()
            
            # Nettoyer le texte des caractères problématiques
            clean_text = ""
            for char in text:
                if ord(char) < 65536:  # Filtrer les caractères Unicode valides
                    clean_text += char
                    
            return clean_text
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte du PDF: {str(e)}")
        logger.error(traceback.format_exc())
        raise Exception(f"Erreur lors de l'extraction du texte du PDF: {str(e)}")

def convert_docx_to_pdf(input_path, output_path):
    """
    Convertit un fichier DOCX en PDF
    """
    try:
        logger.info(f"Conversion du fichier DOCX en PDF: {input_path} -> {output_path}")
        # Méthode 1: Utilisation de docx2pdf (nécessite MS Word sur Windows)
        try:
            from docx2pdf import convert
            convert(input_path, output_path)
            return
        except Exception as e:
            logger.warning(f"Échec de la conversion avec docx2pdf: {str(e)}")
        
        # Méthode 2: Utilisation de LibreOffice (si disponible)
        if os.name == "nt":
            libreoffice_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
            ]
            
            for lo_path in libreoffice_paths:
                if os.path.exists(lo_path):
                    cmd = [
                        lo_path,
                        '--headless',
                        '--convert-to',
                        'pdf',
                        '--outdir',
                        os.path.dirname(output_path),
                        input_path
                    ]
                    
                    process = subprocess.run(cmd, capture_output=True, text=True)
                    
                    if process.returncode == 0:
                        # Renommer le fichier de sortie si nécessaire
                        temp_output = os.path.join(
                            os.path.dirname(output_path),
                            os.path.basename(input_path).replace('.docx', '.pdf').replace('.doc', '.pdf')
                        )
                        
                        if os.path.exists(temp_output) and temp_output != output_path:
                            os.rename(temp_output, output_path)
                        return
        
        # Méthode 3: Utilisation de python-docx et reportlab (solution de secours)
        try:
            logger.info("Tentative de conversion avec python-docx et reportlab")
            from docx import Document
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            
            # Lire le document DOCX
            doc = Document(input_path)
            
            # Créer un document PDF
            pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # Créer un style personnalisé pour les paragraphes
            normal_style = styles["Normal"]
            heading1_style = ParagraphStyle(
                'Heading1',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=12
            )
            heading2_style = ParagraphStyle(
                'Heading2',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=10
            )
            
            # Préparer le contenu du PDF
            content = []
            
            # Parcourir les paragraphes du document DOCX
            for para in doc.paragraphs:
                if para.text.strip():
                    # Déterminer le style en fonction du niveau de titre
                    if para.style.name.startswith('Heading 1') or para.style.name.startswith('Titre 1'):
                        content.append(Paragraph(para.text, heading1_style))
                    elif para.style.name.startswith('Heading 2') or para.style.name.startswith('Titre 2'):
                        content.append(Paragraph(para.text, heading2_style))
                    else:
                        content.append(Paragraph(para.text, normal_style))
                    
                    content.append(Spacer(1, 0.1 * inch))
            
            # Construire le PDF
            pdf_doc.build(content)
            logger.info("Conversion réussie avec python-docx et reportlab")
            return
        except Exception as e:
            logger.warning(f"Échec de la conversion avec python-docx et reportlab: {str(e)}")
        
        raise Exception("Aucune méthode de conversion disponible. Veuillez installer MS Word ou LibreOffice.")
    
    except Exception as e:
        logger.error(f"Erreur lors de la conversion DOCX vers PDF: {str(e)}")
        logger.error(traceback.format_exc())
        raise Exception(f"Erreur lors de la conversion DOCX vers PDF: {str(e)}")

def convert_pdf_to_docx(input_path, output_path):
    """
    Convertit un fichier PDF en DOCX avec une méthode très simple et robuste
    """
    try:
        logger.info(f"Conversion du fichier PDF en DOCX: {input_path} -> {output_path}")
        
        # Créer un nouveau document DOCX
        doc = Document()
        
        # Ajouter un titre au document
        doc.add_heading("Document converti depuis PDF", level=1)
        
        # Extraire le texte avec PyMuPDF (fitz)
        try:
            logger.info("Extraction du texte avec PyMuPDF")
            text = extract_text_from_pdf(input_path)
            
            if text and text.strip():
                # Diviser le texte en paragraphes
                paragraphs = text.split('\n\n')
                
                for i, paragraph in enumerate(paragraphs):
                    if paragraph.strip():
                        # Nettoyer le paragraphe des caractères problématiques pour XML
                        safe_paragraph = clean_text_for_docx(paragraph)
                        
                        # Ajouter le paragraphe au document
                        doc.add_paragraph(safe_paragraph)
                
                logger.info(f"Texte extrait et ajouté au document ({len(paragraphs)} paragraphes)")
            else:
                doc.add_paragraph("Aucun texte n'a pu être extrait du document PDF.")
                logger.warning("Aucun texte extrait du PDF")
        
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction du texte: {str(e)}")
            doc.add_paragraph(f"Erreur lors de l'extraction du texte: {str(e)}")
        
        # Sauvegarder le document
        doc.save(output_path)
        logger.info(f"Document DOCX créé avec succès: {output_path}")
        
    except Exception as e:
        logger.error(f"Erreur lors de la conversion PDF vers DOCX: {str(e)}")
        logger.error(traceback.format_exc())
        raise Exception(f"Erreur lors de la conversion PDF vers DOCX: {str(e)}")

def convert_pdf_to_images(input_path, output_dir):
    """
    Convertit un fichier PDF en images (une image par page) en utilisant PyMuPDF
    
    Args:
        input_path (str): Chemin du fichier PDF à convertir
        output_dir (str): Répertoire de sortie pour les images
        
    Returns:
        tuple: (bool, str, list) indiquant le succès ou l'échec, un message, et la liste des chemins d'images
    """
    try:
        logger.info(f"Conversion du fichier PDF en images: {input_path}")
        
        # Créer le répertoire de sortie si nécessaire
        os.makedirs(output_dir, exist_ok=True)
        
        # Ouvrir le document PDF avec PyMuPDF
        image_paths = []
        with fitz.open(input_path) as pdf:
            # Parcourir chaque page
            for page_num, page in enumerate(pdf):
                # Rendre la page en image avec une résolution plus élevée
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                
                # Définir le chemin de sortie pour l'image
                image_path = os.path.join(output_dir, f"page_{page_num + 1}.png")
                
                # Sauvegarder l'image
                pix.save(image_path)
                
                # Ajouter le chemin à la liste
                image_paths.append(image_path)
                
                logger.info(f"Page {page_num + 1} convertie en image: {image_path}")
        
        return True, f"{len(image_paths)} pages converties en images", image_paths
        
    except Exception as e:
        logger.error(f"Erreur lors de la conversion PDF en images: {str(e)}")
        logger.error(traceback.format_exc())
        return False, f"Erreur lors de la conversion PDF en images: {str(e)}", []

def clean_text_for_docx(text):
    """
    Nettoie le texte pour éliminer les caractères qui posent problème lors de la création de documents DOCX
    """
    if not text:
        return ""
    
    # Nettoyer le texte des caractères problématiques pour XML
    safe_text = ""
    for char in text:
        # Vérifier si le caractère est valide pour XML
        # XML n'accepte que certains caractères de contrôle: tab, CR, LF
        if char in ['\t', '\r', '\n'] or (ord(char) >= 32 and ord(char) != 0xFFFE and ord(char) != 0xFFFF):
            safe_text += char
        else:
            # Remplacer les caractères problématiques
            safe_text += " "
    
    return safe_text

def convert_image_to_text(image_path):
    """
    Extrait le texte d'une image en utilisant OCR
    """
    try:
        # Cette fonction nécessite l'installation de pytesseract et tesseract-ocr
        import pytesseract
        from PIL import Image
        
        logger.info(f"Extraction du texte de l'image: {image_path}")
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img, lang='fra')
        return text
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte de l'image: {str(e)}")
        raise Exception(f"Erreur lors de l'extraction du texte de l'image: {str(e)}")

def extract_text_from_xlsx(file_path):
    """
    Extrait le texte d'un fichier Excel (XLSX)
    
    Args:
        file_path (str): Chemin vers le fichier Excel
        
    Returns:
        str: Texte extrait du fichier
        
    Raises:
        Exception: En cas d'erreur lors de l'extraction
    """
    try:
        logger.info(f"Extraction du texte du fichier Excel (XLSX): {file_path}")
        
        # Utiliser pandas pour lire le fichier Excel
        all_text = []
        
        # Lire toutes les feuilles du fichier Excel
        xlsx = pd.ExcelFile(file_path)
        sheet_names = xlsx.sheet_names
        
        for sheet_name in sheet_names:
            logger.info(f"Lecture de la feuille: {sheet_name}")
            
            # Lire la feuille
            df = pd.read_excel(xlsx, sheet_name=sheet_name)
            
            # Ajouter le nom de la feuille
            all_text.append(f"=== Feuille: {sheet_name} ===")
            
            # Vérifier si le DataFrame est vide
            if df.empty:
                all_text.append("(Feuille vide)")
                continue
            
            # Convertir le DataFrame en texte
            try:
                # Méthode 1: Utiliser to_string pour un format tabulaire
                text = df.to_string(index=False)
                all_text.append(text)
            except Exception as e1:
                logger.warning(f"Erreur lors de la conversion en texte tabulaire: {str(e1)}")
                
                try:
                    # Méthode 2: Convertir en CSV comme alternative
                    csv_buffer = StringIO()
                    df.to_csv(csv_buffer, index=False)
                    text = csv_buffer.getvalue()
                    all_text.append(text)
                except Exception as e2:
                    logger.warning(f"Erreur lors de la conversion en CSV: {str(e2)}")
                    
                    # Méthode 3: Parcourir manuellement les cellules
                    rows_text = []
                    for _, row in df.iterrows():
                        row_values = [str(val) for val in row.values]
                        rows_text.append("\t".join(row_values))
                    
                    if rows_text:
                        all_text.append("\n".join(rows_text))
                    else:
                        all_text.append("(Impossible d'extraire le contenu)")
        
        # Joindre tout le texte
        full_text = "\n\n".join(all_text)
        
        # Nettoyer le texte des caractères problématiques
        clean_text = ""
        for char in full_text:
            if ord(char) < 65536:  # Filtrer les caractères Unicode valides
                clean_text += char
        
        return clean_text
        
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte Excel (XLSX): {str(e)}")
        logger.error(traceback.format_exc())
        raise Exception(f"Erreur lors de l'extraction du texte Excel (XLSX): {str(e)}")

def extract_text_from_xls(file_path):
    """
    Extrait le texte d'un fichier Excel (XLS)
    
    Args:
        file_path (str): Chemin vers le fichier Excel
        
    Returns:
        str: Texte extrait du fichier
        
    Raises:
        Exception: En cas d'erreur lors de l'extraction
    """
    try:
        logger.info(f"Extraction du texte du fichier Excel (XLS): {file_path}")
        
        # Ouvrir le fichier XLS avec xlrd
        workbook = xlrd.open_workbook(file_path)
        
        # Récupérer les noms des feuilles
        sheet_names = workbook.sheet_names()
        
        all_text = []
        
        for sheet_name in sheet_names:
            logger.info(f"Lecture de la feuille: {sheet_name}")
            
            # Ouvrir la feuille
            sheet = workbook.sheet_by_name(sheet_name)
            
            # Ajouter le nom de la feuille
            all_text.append(f"=== Feuille: {sheet_name} ===")
            
            # Vérifier si la feuille est vide
            if sheet.nrows == 0:
                all_text.append("(Feuille vide)")
                continue
            
            # Extraire les données
            sheet_data = []
            
            # Récupérer les en-têtes (première ligne)
            if sheet.nrows > 0:
                headers = [str(sheet.cell_value(0, col_idx)) for col_idx in range(sheet.ncols)]
                sheet_data.append("\t".join(headers))
            
            # Récupérer les données (lignes suivantes)
            for row_idx in range(1, sheet.nrows):
                row_values = [str(sheet.cell_value(row_idx, col_idx)) for col_idx in range(sheet.ncols)]
                sheet_data.append("\t".join(row_values))
            
            # Ajouter les données de la feuille au texte
            all_text.append("\n".join(sheet_data))
        
        # Joindre tout le texte
        full_text = "\n\n".join(all_text)
        
        # Nettoyer le texte des caractères problématiques
        clean_text = ""
        for char in full_text:
            if ord(char) < 65536:  # Filtrer les caractères Unicode valides
                clean_text += char
        
        return clean_text
        
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte Excel (XLS): {str(e)}")
        logger.error(traceback.format_exc())
        raise Exception(f"Erreur lors de l'extraction du texte Excel (XLS): {str(e)}")

def extract_text_from_file(file_path):
    """
    Extrait le texte d'un fichier en détectant automatiquement son type (PDF, DOCX/DOC, Excel, ou image)
    
    Args:
        file_path (str): Chemin vers le fichier
        
    Returns:
        str: Texte extrait du fichier
        
    Raises:
        Exception: Si le format de fichier n'est pas pris en charge ou en cas d'erreur
    """
    try:
        logger.info(f"Extraction du texte du fichier: {file_path}")
        
        # Déterminer l'extension du fichier
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # Extraire le texte en fonction du type de fichier
        if file_extension in ['.pdf']:
            return extract_text_from_pdf(file_path)
        
        elif file_extension in ['.docx', '.doc']:
            return extract_text_from_docx(file_path)
        
        elif file_extension in ['.xlsx', '.xls']:
            if file_extension == '.xlsx':
                return extract_text_from_xlsx(file_path)
            else:
                return extract_text_from_xls(file_path)
        
        elif file_extension in ['.csv']:
            return extract_text_from_csv(file_path)
        
        elif file_extension in ['.txt']:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
        
        elif file_extension in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.tif']:
            return convert_image_to_text(file_path)
        
        else:
            raise Exception(f"Format de fichier non pris en charge: {file_extension}")
    
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte: {str(e)}")
        logger.error(traceback.format_exc())
        raise Exception(f"Erreur lors de l'extraction du texte: {str(e)}")

def extract_text_from_csv(file_path):
    """
    Extrait le texte d'un fichier CSV
    
    Args:
        file_path (str): Chemin vers le fichier CSV
        
    Returns:
        str: Texte extrait du fichier
        
    Raises:
        Exception: En cas d'erreur lors de l'extraction
    """
    try:
        logger.info(f"Extraction du texte du fichier CSV: {file_path}")
        
        # Détecter l'encodage du fichier
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            result = chardet.detect(raw_data)
            encoding = result['encoding']
        
        logger.info(f"Encodage détecté: {encoding}")
        
        # Détecter le délimiteur
        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
            sample = f.read(4096)
            
            # Compter les occurrences des délimiteurs courants
            delimiters = [',', ';', '\t', '|']
            counts = {d: sample.count(d) for d in delimiters}
            
            # Choisir le délimiteur le plus fréquent
            delimiter = max(counts, key=counts.get)
            
            if counts[delimiter] == 0:
                # Aucun délimiteur trouvé, utiliser la virgule par défaut
                delimiter = ','
        
        logger.info(f"Délimiteur détecté: '{delimiter}'")
        
        # Lire le fichier CSV
        df = pd.read_csv(file_path, delimiter=delimiter, encoding=encoding, on_bad_lines='warn')
        
        # Convertir en texte
        text = df.to_string(index=False)
        
        return text
        
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte CSV: {str(e)}")
        logger.error(traceback.format_exc())
        raise Exception(f"Erreur lors de l'extraction du texte CSV: {str(e)}")

def convert_text_to_csv(text, output_path, delimiter=',', has_header=True):
    """
    Convertit du texte en fichier CSV
    
    Args:
        text (str): Texte à convertir
        output_path (str): Chemin de sortie pour le fichier CSV
        delimiter (str): Délimiteur à utiliser
        has_header (bool): Si True, la première ligne est considérée comme l'en-tête
        
    Returns:
        bool: True si la conversion a réussi, False sinon
        
    Raises:
        Exception: En cas d'erreur lors de la conversion
    """
    try:
        logger.info(f"Conversion de texte en CSV avec délimiteur '{delimiter}'")
        
        # Diviser le texte en lignes
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        
        if not lines:
            raise ValueError("Le texte ne contient aucune ligne valide")
        
        # Créer le répertoire de sortie si nécessaire
        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
        
        # Déterminer le délimiteur d'entrée pour l'analyse
        input_delimiter = None
        if delimiter == "auto":
            # Détecter automatiquement le délimiteur
            first_line = lines[0]
            if '\t' in first_line:
                input_delimiter = '\t'
            elif ';' in first_line:
                input_delimiter = ';'
            elif ',' in first_line:
                input_delimiter = ','
            else:
                input_delimiter = ' '
            logger.info(f"Délimiteur auto-détecté: '{input_delimiter}'")
        elif delimiter == "space":
            input_delimiter = ' '
        elif delimiter == "tab" or delimiter == "\\t":
            input_delimiter = '\t'
        else:
            input_delimiter = delimiter
            
        logger.info(f"Délimiteur utilisé pour l'analyse: '{input_delimiter}'")
        
        # Déterminer le délimiteur de sortie pour le CSV
        output_delimiter = ',' if input_delimiter == ' ' else input_delimiter
        
        # Écrire le fichier CSV
        with open(output_path, 'w', newline='', encoding='utf-8') as csv_file:
            writer = csv.writer(csv_file, delimiter=output_delimiter)
            
            for i, line in enumerate(lines):
                # Traiter la ligne en fonction du délimiteur
                if input_delimiter == ' ':
                    # Utiliser une expression régulière pour diviser sur les espaces multiples
                    row = re.split(r'\s+', line.strip())
                else:
                    row = line.split(input_delimiter)
                
                # Nettoyer les valeurs (supprimer les espaces en début/fin)
                row = [cell.strip() for cell in row]
                
                writer.writerow(row)
        
        logger.info(f"Conversion réussie, fichier CSV créé: {output_path}")
        return True
    
    except Exception as e:
        logger.error(f"Erreur lors de la conversion de texte en CSV: {str(e)}")
        logger.error(traceback.format_exc())
        raise Exception(f"Erreur lors de la conversion de texte en CSV: {str(e)}")

def convert_text_to_csv_interactive(text, output_path, delimiter=',', has_header=True):
    """
    Convertit du texte en fichier CSV avec des options interactives
    
    Args:
        text (str): Texte à convertir
        output_path (str): Chemin de sortie pour le fichier CSV
        delimiter (str): Délimiteur à utiliser
        has_header (bool): Si True, la première ligne est considérée comme l'en-tête
        
    Returns:
        bool: True si la conversion a réussi, False sinon
        
    Raises:
        Exception: En cas d'erreur lors de la conversion
    """
    try:
        logger.info(f"Conversion interactive de texte en CSV avec délimiteur '{delimiter}'")
        
        # Diviser le texte en lignes
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        
        if not lines:
            raise ValueError("Le texte ne contient aucune ligne valide")
        
        # Créer le répertoire de sortie si nécessaire
        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
        
        # Déterminer le délimiteur d'entrée pour l'analyse
        input_delimiter = None
        if delimiter == "auto":
            # Détecter automatiquement le délimiteur
            first_line = lines[0]
            if '\t' in first_line:
                input_delimiter = '\t'
            elif ';' in first_line:
                input_delimiter = ';'
            elif ',' in first_line:
                input_delimiter = ','
            else:
                input_delimiter = ' '
            logger.info(f"Délimiteur auto-détecté: '{input_delimiter}'")
        elif delimiter == "space":
            input_delimiter = ' '
        elif delimiter == "tab" or delimiter == "\\t":
            input_delimiter = '\t'
        else:
            input_delimiter = delimiter
            
        logger.info(f"Délimiteur utilisé pour l'analyse: '{input_delimiter}'")
        
        # Déterminer le délimiteur de sortie pour le CSV
        output_delimiter = ',' if input_delimiter == ' ' else input_delimiter
        
        # Écrire le fichier CSV
        with open(output_path, 'w', newline='', encoding='utf-8') as csv_file:
            writer = csv.writer(csv_file, delimiter=output_delimiter)
            
            for i, line in enumerate(lines):
                # Traiter la ligne en fonction du délimiteur
                if input_delimiter == ' ':
                    # Utiliser une expression régulière pour diviser sur les espaces multiples
                    row = re.split(r'\s+', line.strip())
                else:
                    row = line.split(input_delimiter)
                
                # Nettoyer les valeurs (supprimer les espaces en début/fin)
                row = [cell.strip() for cell in row]
                
                writer.writerow(row)
        
        logger.info(f"Conversion réussie, fichier CSV créé: {output_path}")
        return True
    
    except Exception as e:
        logger.error(f"Erreur lors de la conversion de texte en CSV: {str(e)}")
        logger.error(traceback.format_exc())
        raise Exception(f"Erreur lors de la conversion de texte en CSV: {str(e)}")
